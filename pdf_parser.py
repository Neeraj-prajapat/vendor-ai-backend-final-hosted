# https://chatgpt.com/share/6817b643-5740-8007-81a2-714b9707e35a
#Yes — ✅ pdfplumber can extract both text and tables from PDFs, and that's one of its core strengths.
#? Explanation of this chat:- https://chatgpt.com/share/681975a1-f7bc-8012-bf7e-690074978987 

import re
from typing import List
import fitz  # PyMuPDF
import pdfplumber
from docx import Document

# -------------------------------
# TABLE PROCESSING UTILITIES
# -------------------------------

def clean_table(table: List[List[str]]) -> List[List[str]]:
    """Remove empty rows/columns and clean cell content"""
    cleaned = [
        [cell.strip() if cell else "" for cell in row]
        for row in table if any((cell.strip() if cell else "") for cell in row)
    ]
    
    if cleaned:
        transposed = list(zip(*cleaned))
        kept_cols = [col for col in transposed if any(cell for cell in col)]
        cleaned = list(zip(*kept_cols))
        cleaned = [list(row) for row in cleaned]
    return cleaned

def standardize_columns(table: List[List[str]]) -> List[List[str]]:
    """Ensure all rows have same number of columns"""
    if not table: return []
    max_cols = max(len(row) for row in table)
    return [row + [""] * (max_cols - len(row)) for row in table]

def merge_wrapped_rows(table: List[List[str]]) -> List[List[str]]:
    """Merge rows where content is wrapped across multiple lines"""
    merged = []
    previous_row = None
    
    for row in table:
        if previous_row:
            if not row[0].strip() and any(cell.strip() for cell in row[1:]):
                for i in range(len(previous_row)):
                    previous_row[i] += " " + row[i].strip()
            else:
                merged.append(previous_row)
                previous_row = row
        else:
            previous_row = row
    if previous_row:
        merged.append(previous_row)
    return merged

def table_to_markdown(table: List[List[str]]) -> str:
    """Convert cleaned table to markdown format"""
    if not table or len(table) < 2: return ""
    
    header = table[0]
    separator = ["---"] * len(header)
    body = table[1:]
    
    markdown = [
        f"| {' | '.join(header)} |",
        f"| {' | '.join(separator)} |"
    ]
    
    for row in body:
        markdown.append(f"| {' | '.join(row)} |")
    
    return '\n'.join(markdown)

# -------------------------------
# TEXT EXTRACTION METHODS
# -------------------------------

def extract_text_pdfplumber(file_path: str) -> str:
    """Extract content with optimized table processing for LLMs"""
    content = []
    
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            # Extract text
            text = page.extract_text()
            if text:
                content.append(f"\n\n=== PAGE {page_num} TEXT ===")
                content.extend(line.strip() for line in text.split('\n') if line.strip())
            
            # Process tables
            tables = page.extract_tables({
                "vertical_strategy": "text",
                "horizontal_strategy": "text",
                "join_tolerance": 15
            })
            
            for table_num, table in enumerate(tables, 1):
                cleaned = clean_table(table)
                if not cleaned: continue
                
                processed = merge_wrapped_rows(
                    standardize_columns(cleaned)
                )
                
                if len(processed) >= 2:
                    content.append(f"\n\n=== PAGE {page_num} TABLE {table_num} ===")
                    content.append(table_to_markdown(processed))
            
            content.append(f"\n{'='*40} END PAGE {page_num} {'='*40}")
    
    return '\n'.join(content)

def extract_text_pymupdf(file_path: str) -> str:
    """Extract cleaned text using PyMuPDF"""
    doc = fitz.open(file_path)
    return '\n'.join(
        line.strip() for page in doc
        for line in page.get_text("text").split('\n')
        if line.strip()
    )

def extract_text_docx(file_path: str) -> str:
    """Extract DOCX content with table support"""
    doc = Document(file_path)
    content = []
    
    # Paragraphs
    content.extend(
        para.text.strip() 
        for para in doc.paragraphs 
        if para.text.strip()
    )
    
    # Tables
    for table in doc.tables:
        processed = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                processed.append(cells)
        
        if len(processed) >= 2:
            content.append('\n' + table_to_markdown(processed))
    
    return '\n'.join(content)

def extract_text(file_path: str, method: str = "pdfplumber") -> str:
    """
    Unified text extraction with LLM-optimized table processing
    - PDF: 'pdfplumber' for tables, 'pymupdf' for pure text
    - DOCX: Structured text + markdown tables
    """
    lower_path = file_path.lower()
    
    if lower_path.endswith(".docx"):
        return extract_text_docx(file_path)
    elif lower_path.endswith(".pdf"):
        if method == "pymupdf":
            return extract_text_pymupdf(file_path)
        return extract_text_pdfplumber(file_path)
    raise ValueError("Unsupported file format. Only .pdf and .docx are supported.")

